import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = create_element('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def convert_md_to_docx():
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    PRIMARY = RGBColor(0x1E, 0x3A, 0x8A)
    SECONDARY = RGBColor(0x02, 0x84, 0xC7)
    DARK_TEXT = RGBColor(0x1E, 0x29, 0x3B)

    with open("d:/alba/Akademik/Semester 6/Pengujian Perangkat Lunak/Pertemuan 15/Laporan_Pengujian_PPL_Login_Register.md", "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []

    for line in lines:
        line_str = line.rstrip('\n')

        # Code block toggle
        if line_str.startswith("```"):
            if in_code_block:
                # End code block
                in_code_block = False
                code_text = "\n".join(code_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "0F172A")
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)
                doc.add_paragraph() # spacer
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line_str)
            continue

        if line_str.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_str[2:])
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = PRIMARY
        elif line_str.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = PRIMARY
            p.paragraph_format.space_before = Pt(14)
        elif line_str.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line_str[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = SECONDARY
            p.paragraph_format.space_before = Pt(10)
        elif line_str.startswith("- ") or line_str.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line_str[2:])
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT
        elif line_str.startswith("|"):
            # Table handling simplified / skip header divider
            if "---|---" in line_str or "|---" in line_str:
                continue
            cols = [c.strip() for c in line_str.split("|")[1:-1]]
            if len(cols) > 1:
                # Append row
                pass # Handled as text or basic table if needed
        elif line_str.strip() == "---":
            doc.add_paragraph()
        elif line_str.strip():
            p = doc.add_paragraph()
            run = p.add_run(line_str)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT

    docx_path = "d:/alba/Akademik/Semester 6/Pengujian Perangkat Lunak/Pertemuan 15/Laporan_Pengujian_PPL_Login_Register.docx"
    doc.save(docx_path)
    print("DOCX successfully created:", docx_path)

if __name__ == "__main__":
    convert_md_to_docx()
